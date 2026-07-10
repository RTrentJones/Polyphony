"""Assemble a book's current draft state into Markdown / DOCX / EPUB.

Exports use each scene's editable `content` (falling back to
`generated_content`), never raw generation output when an edit exists.
"""

import html
import io
from dataclasses import dataclass, field


@dataclass
class SceneExport:
    title: str
    content: str


@dataclass
class ChapterExport:
    title: str
    summary: str = ""
    scenes: list[SceneExport] = field(default_factory=list)


@dataclass
class BookExport:
    title: str
    author: str = ""
    synopsis: str = ""
    chapters: list[ChapterExport] = field(default_factory=list)


def scene_text(scene) -> str:
    """The draft state: edited content wins over raw generation output."""
    return (scene.content or scene.generated_content or "").strip()


def to_markdown(book: BookExport) -> str:
    parts = [f"# {book.title}"]
    if book.author:
        parts.append(f"*by {book.author}*")
    if book.synopsis:
        parts.append(f"> {book.synopsis}")
    parts.append("")
    for i, chapter in enumerate(book.chapters, 1):
        parts.append(f"## Chapter {i}: {chapter.title}")
        parts.append("")
        for scene in chapter.scenes:
            if scene.content:
                parts.append(scene.content)
                parts.append("")
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def to_docx(book: BookExport) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(book.title, level=0)
    if book.author:
        doc.add_paragraph(f"by {book.author}")
    if book.synopsis:
        doc.add_paragraph(book.synopsis)
    for i, chapter in enumerate(book.chapters, 1):
        doc.add_heading(f"Chapter {i}: {chapter.title}", level=1)
        for scene in chapter.scenes:
            for para in scene.content.split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def to_epub(book: BookExport) -> bytes:
    from ebooklib import epub

    e_book = epub.EpubBook()
    e_book.set_identifier(f"polyphony-{book.title[:40]}")
    e_book.set_title(book.title)
    e_book.set_language("en")
    if book.author:
        e_book.add_author(book.author)

    spine: list = ["nav"]
    toc = []
    for i, chapter in enumerate(book.chapters, 1):
        chapter_title = f"Chapter {i}: {chapter.title}"
        body = [f"<h1>{html.escape(chapter_title)}</h1>"]
        for scene in chapter.scenes:
            for para in scene.content.split("\n\n"):
                para = para.strip()
                if para:
                    body.append(f"<p>{html.escape(para)}</p>")
        item = epub.EpubHtml(
            title=chapter_title,
            file_name=f"chapter_{i}.xhtml",
            lang="en",
        )
        item.content = "\n".join(body)
        e_book.add_item(item)
        spine.append(item)
        toc.append(item)

    e_book.toc = toc
    e_book.spine = spine
    e_book.add_item(epub.EpubNcx())
    e_book.add_item(epub.EpubNav())

    buffer = io.BytesIO()
    epub.write_epub(buffer, e_book)
    return buffer.getvalue()


CONTENT_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "epub": "application/epub+zip",
}
